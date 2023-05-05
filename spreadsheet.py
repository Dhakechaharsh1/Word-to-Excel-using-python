import os
import glob
from docx import Document
import openpyxl
from openpyxl import Workbook

def process_invoices():
    # Create a new Excel workbook and add a sheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Invoices"
    
    # Add headers to the sheet
    headers = ["Invoice ID", "Total Number of Products Purchased", "Subtotal", "Tax", "Total"]
    for col_num, header in enumerate(headers, 1):
        ws.cell(row=1, column=col_num, value=header)

    # Process all Word documents
    for file in glob.glob("*.docx"):
        doc_file = open(file, "rb")
        doc = Document(doc_file)
        invoice_id = doc.paragraphs[0].text.replace("INV", "")

        # Ignore the "PRODUCTS" line and the last empty line
        product_lines = doc.paragraphs[1].text.split("\n")[1:-1]  
        total_products = sum(int(line.split(":")[1]) for line in product_lines)
        subtotal = float(doc.paragraphs[2].text.split(":")[1].split("\n")[0])
        tax = float(doc.paragraphs[2].text.split(":")[2].split("\n")[0])
        total = float(doc.paragraphs[2].text.split(":")[3])

        # Write the extracted data to the sheet
        row = [invoice_id, total_products, subtotal, tax, total]
        ws.append(row)

    # Save the workbook
    wb.save("New Spreadsheet.xlsx")

if __name__ == "__main__":
    process_invoices()
